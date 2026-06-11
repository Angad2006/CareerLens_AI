"""
Resume file parser — extracts text from PDF and DOCX files.
"""
import os
from pathlib import Path


def extract_text_from_pdf(file_path: str) -> dict:
    """
    Extract text from a PDF file using pdfplumber.
    Returns dict with text content and page count.
    """
    import pdfplumber

    text_parts = []
    page_count = 0

    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    return {
        "text": "\n\n".join(text_parts),
        "page_count": page_count
    }


def extract_text_from_docx(file_path: str) -> dict:
    """
    Extract text from a DOCX file using python-docx.
    Returns dict with text content and estimated page count.
    """
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    full_text = "\n".join(text_parts)

    # Estimate page count (~400 words per page)
    word_count = len(full_text.split())
    estimated_pages = max(1, word_count // 400)

    return {
        "text": full_text,
        "page_count": estimated_pages
    }


def parse_resume(file_path: str) -> dict:
    """
    Parse a resume file and extract text. Auto-detects format.
    Returns dict with text, page_count, and format.
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        result = extract_text_from_pdf(file_path)
        result["format"] = "pdf"
    elif ext == ".docx":
        result = extract_text_from_docx(file_path)
        result["format"] = "docx"
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    if not result["text"].strip():
        raise ValueError("Could not extract any text from the resume. The file may be image-based or corrupted.")

    return result
